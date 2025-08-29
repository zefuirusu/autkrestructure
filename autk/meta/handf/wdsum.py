#!/usr/bin/env python
#

import re
import os
from threading import Thread,Lock
from docx import Document

from autk.gentk.funcs import regex_filter,start_thread_list
from autk.meta.handf.findfile import find_regex


class WordSum:
    def __init__(
        self,
        path:str,
    ):
        self.path=path
        self.title=self.path.split(os.sep)[-1]
        self.matched_paragraphs=[]
        pass
    def single_search(self,item):
        matches=regex_filter(
            item,
            map(
                lambda p:p.text,
                Document(self.path).paragraphs
            ),
            match_mode=False
        )
        for single_match in matches:
            if single_match is not None:
                self.matched_paragraphs.append(
                    single_match
                )
            else:
                pass
        pass
    def save(self,save_path):
        if len(self.matched_paragraphs)==0:
            pass
        else:
            if os.path.isfile(save_path):
                target_docx=Document(save_path)
            else:
                target_docx=Document()
            target_docx.add_heading(
                'Summary from:{}'.format(
                    self.title.join(['《','》'])
                ),
                level=2
            )
            target_docx.add_heading(
                'source:{}'.format(self.path),
                level=3
            )
            for p in self.matched_paragraphs:
                target_docx.add_paragraph(p)
            target_docx.save(save_path)
            print(
                'summary from `{}` saved to:{}'.format(
                    self.title,
                    save_path
                )
            )
        pass
    pass
def docxSum(item:str,sdir:str,save_path:str):
    if os.path.isfile(save_path):
        pass
    else:
        blank_docx=Document()
        blank_docx.save(save_path)
    file_list=find_regex(
        r'^[^\~\$].*\.docx$',
        sdir,
        match=True
    )[0]
    result_wd_collection=[]
    def __single_search(docx_path):
        sch_lock=Lock()
        sch_lock.acquire()
        ws=WordSum(docx_path)
        ws.single_search(item)
        if len(ws.matched_paragraphs)>0:
            print(
                'search result from 《{}》:\n\t{}'.format(
                    ws.title,
                    ws.matched_paragraphs
                )
            )
        ws.save(save_path)
        result_wd_collection.append(ws)
        sch_lock.release()
        pass
    thli=[]
    for docx_path in file_list:
        thli.append(
            Thread(
                target=__single_search,
                args=(docx_path,),
                name=__single_search.__name__+docx_path
            )
        )
        continue
    start_thread_list(thli)
    for wd in result_wd_collection:
        wd.save(save_path)
    pass
if __name__=='__main__':
    pass
